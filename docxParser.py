import os
from pathlib import Path

import docx.document
import xml.etree.ElementTree as ET
from docx.api import Document
from docx.document import Document as doctwo
from docx.shared import Pt, Cm, Emu
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT as WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENTATION
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from io import StringIO
import re
# from docx.enum.text import WD_ALIGN_PARAGRAPH



regex_transform = {
    "Таблица <N> - <Название>.": "^Таблица [1-9]*(.[1-9])* - .*$",
    "<Название>.": ".*",
    "Рисунок <N> - <Название>.": "^Рисунок [1-9]*(.[1-9])* - .*$"
}


text_checklist = {
    "font_name": "Times New Roman",
    "font_size": 13.0,
    "font_bald": False,
    "font_italic": False,
    "font_underline": False,
    "font_color": None,
    "font_back_color": None,
    "alignment": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
    "keep_with_next": False,
    "page_break_before": False,
    "space_before": 0.0,
    "space_after": 0.0,
    "left_indent": 0,
    "right_indent": 0,
    "first_line_indent": 1.25,
    "line_spacing": 1.5,
}
heading1_checklist = {
    "font_name": "Times New Roman",
    "font_size": 16.0,
    "font_bald": True,
    "font_italic": False,
    "font_underline": False,
    "font_color": None,
    "font_back_color": None,
    "alignment": WD_PARAGRAPH_ALIGNMENT.CENTER,
    "keep_with_next": True,
    "page_break_before": True,
    "space_before": 0.0,
    "space_after": 12.0,
    "left_indent": 0,
    "right_indent": 0,
    "first_line_indent": 0.0,
    "line_spacing": 1.0,
    "is_list": True
}
heading2_checklist = {
    "font_name": "Times New Roman",
    "font_size": 14.0,
    "font_bald": True,
    "font_italic": False,
    "font_underline": False,
    "font_color": None,
    "font_back_color": None,
    "alignment": WD_PARAGRAPH_ALIGNMENT.CENTER,
    "keep_with_next": True,
    "page_break_before": False,
    "space_before": 12.0,
    "space_after": 6.0,
    "left_indent": 0,
    "right_indent": 0,
    "first_line_indent": 0.0,
    "line_spacing": 1.0,
    "is_list": True
}
heading3_checklist = {
    "font_name": "Times New Roman",
    "font_size": 13.0,
    "font_bald": True,
    "font_italic": False,
    "font_underline": False,
    "font_color": None,
    "font_back_color": None,
    "alignment": WD_PARAGRAPH_ALIGNMENT.CENTER,
    "keep_with_next": True,
    "page_break_before": False,
    "space_before": 8.0,
    "space_after": 4.0,
    "left_indent": 0,
    "right_indent": 0,
    "first_line_indent": 0.0,
    "line_spacing": 1.0,
    "is_list": True
}
list_checklist = {
    "font_name": "Times New Roman",
    "font_size": 13.0,
    "font_bald": False,
    "font_italic": False,
    "font_underline": False,
    "font_color": None,
    "font_back_color": None,
    "keep_with_next": False,
    "page_break_before": False,
    "space_before": 0.0,
    "space_after": 0.0,
    "line_spacing": 1.5,
}
table_headings_checklist = {
    "font_name": "Times New Roman",
    "font_size": 11.0,
    "font_bald": True,
    "font_italic": False,
    "font_underline": False,
    "font_color": None,
    "font_back_color": None,
    "alignment": WD_PARAGRAPH_ALIGNMENT.CENTER,
    "vert_alignment": WD_ALIGN_VERTICAL.CENTER,
    "keep_with_next": False,
    "page_break_before": False,
    "space_before": 0.0,
    "space_after": 0.0,
    "left_indent": 0,
    "right_indent": 0,
    "first_line_indent": 0.0,
    "line_spacing": 1.0,
}
table_text_checklist = {
    "font_name": "Times New Roman",
    "font_size": 11.0,
    "font_bald": False,
    "font_italic": False,
    "font_underline": False,
    "font_color": None,
    "font_back_color": None,
    "alignment": WD_PARAGRAPH_ALIGNMENT.LEFT,
    "vert_alignment": WD_ALIGN_VERTICAL.CENTER,
    "keep_with_next": False,
    "page_break_before": False,
    "space_before": 0.0,
    "space_after": 0.0,
    "left_indent": 0,
    "right_indent": 0,
    "first_line_indent": 0.0,
    "line_spacing": 1.0,
}
table_name_checklist = {
    "format_regex": "Таблица <N> - <Название>.",
    "font_name": "Times New Roman",
    "font_size": 13.0,
    "font_color": None,
    "font_back_color": None,
    "alignment": WD_PARAGRAPH_ALIGNMENT.LEFT,
    "keep_with_next": False,
    "page_break_before": False,
    "space_before": 13.0,
    "space_after": 0.0,
    "left_indent": 0,
    "right_indent": 0,
    "first_line_indent": 0.0,
    "line_spacing": 1.0,
}
image_checklist = {
    "alignment": WD_PARAGRAPH_ALIGNMENT.CENTER,
    "keep_with_next": True,
    "space_before": 6.0,
    "space_after": 0.0,
    "first_line_indent": 0,
    "line_spacing": 1,
}
image_name_checklist = {
    "format_regex": "Рисунок {N} - <Название>.",
    "font_name": "Times New Roman",
    "font_size": 11.0,
    "font_bald": True,
    "font_italic": True,
    "font_color": None,
    "font_back_color": None,
    "alignment": WD_PARAGRAPH_ALIGNMENT.CENTER,
    "keep_with_next": False,
    "page_break_before": False,
    "space_before": 0.0,
    "space_after": 6.0,
    "left_indent": 0,
    "right_indent": 0,
    "first_line_indent": 0,
    "line_spacing": 1,
}
param_to_comment = {
    "format_regex": "Формат",
    "font_name": "Тип шрифта",
    "font_size": "Размер шрифта",
    "font_bald": "Полужирный",
    "font_italic": "Курсив",
    "font_underline": "Подчёркивание",
    "font_color": "Цвет текста",
    "font_back_color": "Цвет подчёркивания",
    "alignment": "Выравнивание",
    "vert_alignment": "Выравнивание по вертикали",
    "keep_with_next": "Не отрывать от следующего",
    "page_break_before": "С новой страницы",
    "space_before": "Верт отступ перед абзацем",
    "space_after": "Верт отступ после абзаца",
    "left_indent": "Отступ слева",
    "right_indent": "Отступ справа",
    "first_line_indent": "Отступ первой строки",
    "line_spacing": "Межстрочный интервал",
    "is_list": "Список",
    "top_margin": "Верхнее поле",
    "bottom_margin": "Нижнее поле",
    "left_margin": "Левое поле",
    "right_margin": "Правое поле",
    "orientation": "Ориентация",
}
var_to_comment = {
    True: "Да",
    False: "Нет",
    None: "Нет",
}
alignment_to_comment = {
    WD_PARAGRAPH_ALIGNMENT.LEFT: "Левый край",
    WD_PARAGRAPH_ALIGNMENT.CENTER: "По центру",
    WD_PARAGRAPH_ALIGNMENT.RIGHT: "Правый край",
    WD_PARAGRAPH_ALIGNMENT.JUSTIFY: "По ширине",
}
vert_alignment_to_comment = {
    WD_ALIGN_VERTICAL.TOP: "Верхний край",
    WD_ALIGN_VERTICAL.CENTER: "По центру",
    WD_ALIGN_VERTICAL.BOTTOM: "Нижний край",
    WD_ALIGN_VERTICAL.BOTH: "Ребята, не стоит вскрывать эту тему...",
}
orientation_to_comment = {
    WD_ORIENTATION.PORTRAIT: "Книжная",
    WD_ORIENTATION.LANDSCAPE: "Альбомная"
}
margins_check = {
    "top_margin": 2.0,
    "bottom_margin": 2.0,
    "left_margin": 3.0,
    "right_margin": 1.5,
    "orientation": WD_ORIENTATION.PORTRAIT
}

enable_elem_check = {
    "table_headings_top": True,
    "table_headings_left": False,
    "paragraph_before_table": True,
    "enable_pic_title": True
}


def set_settings(text_check, h1_check, h2_check, h3_check, table_name_check, table_heading_check, table_text_check, list_check, page_check, pic_check, pic_name_check):
    if isinstance(text_check, dict):
        for name, parameter in text_check.items():
            if name in text_checklist:
                text_checklist[name] = parameter
    if isinstance(h1_check, dict):
        for name, parameter in h1_check.items():
            if name in heading1_checklist:
                heading1_checklist[name] = parameter
    if isinstance(h2_check, dict):
        for name, parameter in h2_check.items():
            if name in heading2_checklist:
                heading2_checklist[name] = parameter
    if isinstance(h3_check, dict):
        for name, parameter in h3_check.items():
            if name in heading3_checklist:
                heading3_checklist[name] = parameter
    if isinstance(table_name_check, dict):
        for name, parameter in table_name_check.items():
            if name in table_name_checklist:
                table_name_checklist[name] = parameter
    if isinstance(table_heading_check, dict):
        for name, parameter in table_heading_check.items():
            if name in table_headings_checklist:
                table_headings_checklist[name] = parameter
            elif name == "heading_left":
                enable_elem_check["table_headings_left"] = parameter
            elif name == "heading_top":
                enable_elem_check["table_headings_top"] = parameter
    if isinstance(table_text_check, dict):
        for name, parameter in table_text_check.items():
            if name in table_text_checklist:
                table_name_checklist[name] = parameter
    if isinstance(list_check, dict):
        for name, parameter in list_check.items():
            if name in list_checklist:
                list_checklist[name] = parameter
    if isinstance(page_check, dict):
        for name, parameter in page_check.items():
            if name in margins_check:
                margins_check[name] = parameter
    if isinstance(pic_check, dict):
        for name, parameter in pic_check.items():
            if name in image_checklist:
                image_checklist[name] = parameter
    if isinstance(pic_name_check, dict):
        for name, parameter in pic_name_check.items():
            if name in image_name_checklist:
                image_name_checklist[name] = parameter
            elif name == "enable_pic_title":
                enable_elem_check["enable_pic_title"] = parameter


def is_list(paragraph):
    return len(paragraph._element.xpath('./w:pPr/w:numPr')) > 0


def get_error_comment(expected: dict, received: dict):
    comments = set()
    for key, val in expected.items():
        comment = ""
        if key in received.keys():
            # Отдельная обработка формата через RegEx
            if key == 'format_regex':
                regex_check = re.fullmatch(fr'{regex_transform[val]}', received[key])
                if not regex_check:
                    comment += f"{param_to_comment[key]}: "
                    comment += f"{val}; "
                    comment += f"Получено: {received[key]}.\n"
                    comments.add(comment)
            elif val != received[key]:

                comment += f"{param_to_comment[key]}: "
                # У выравниваний значения приравниваются к 0..3, поэтому нужен костыль
                if key == "alignment":
                    comment += f"Ожидалось: {alignment_to_comment[val]}; "
                    comment += f"Получено: {alignment_to_comment[received[key]]}.\n"
                elif key == "vert_alignment":
                    comment += f"Ожидалось: {vert_alignment_to_comment[val]}; "
                    comment += f"Получено: {vert_alignment_to_comment[received[key]]}.\n"
                elif key == "orientation":
                    comment += f"Ожидалось: {orientation_to_comment[val]}; "
                    comment += f"Получено: {orientation_to_comment[received[key]]}.\n"
                else:
                    if val is None or isinstance(val, bool):
                        comment += f"Ожидалось: {var_to_comment[val]}; "
                    else:
                        comment += f"Ожидалось: {val}; "
                    if received[key] is None or isinstance(received[key], bool):
                        comment += f"Получено: {var_to_comment[received[key]]}.\n"
                    else:
                        comment += f"Получено: {received[key]}.\n"
                comments.add(comment)
    return comments


def get_run_properties(document, p, run):
    """ Параметры параграфа """
    st = p.style
    formatting = p.paragraph_format  # Формат параграфа
    st_formatting = p.style.paragraph_format  # Формат параграфа, заданный стилем
    default = st.base_style if st.base_style else document.styles["Normal"] # Стиль по умолчанию
    def_formatting = default.paragraph_format

    # Пришлось всему написать is not None, т.к может встретиться 0
    # Название шрифта
    font_name = run.font.name if run.font.name is not None else \
        st.font.name if st.font.name is not None else \
            default.font.name if default.font.name is not None else "Calibri"
    # Размер шрифта
    font_size = run.font.size.pt if run.font.size is not None else \
        st.font.size.pt if st.font.size is not None else \
            default.font.size.pt if default.font.size is not None else 11
    # Полужирный
    font_bald = run.bold if run.bold is not None else \
        st.font.bold if st.font.bold is not None else \
            default.font.bold if default.font.bold is not None else False
    # Курсив
    font_italic = run.italic if run.italic is not None else \
        st.font.italic if st.font.italic is not None else \
            default.font.italic if default.font.italic is not None else False
    # Подчёркивание
    font_underline = run.underline if run.underline is not None else \
        st.font.underline if st.font.underline is not None else \
            default.font.underline if default.font.underline is not None else False
    # Цвет текста
    font_color = run.font.color.rgb if run.font.color.rgb is not None else \
        st.font.color.rgb if st.font.color.rgb is not None else \
            default.font.color.rgb if default.font.color.rgb is not None else None

    if font_color is not None and font_color[0] == font_color[1] == font_color[2] == 0:
        font_color = None
    # Цвет подчёркивания текста
    font_back_color = run.font.highlight_color if run.font.highlight_color is not None else \
        st.font.highlight_color if st.font.highlight_color is not None else \
            default.font.highlight_color if default.font.highlight_color is not None else None


    # Выравнивание (лево/центр/право/ширина)
    alignment = formatting.alignment if formatting.alignment is not None else \
        st_formatting.alignment if st_formatting.alignment is not None else \
            def_formatting.alignment if def_formatting.alignment is not None else WD_PARAGRAPH_ALIGNMENT.LEFT
    # Не отрывать от следующего
    keep_with_next = formatting.keep_with_next if formatting.keep_with_next is not None else \
        st_formatting.keep_with_next if st_formatting.keep_with_next is not None else \
            def_formatting.keep_with_next if def_formatting.keep_with_next is not None else False
    # С новой страницы
    page_break_before = formatting.page_break_before if formatting.page_break_before is not None else \
        st_formatting.page_break_before if st_formatting.page_break_before is not None else \
            def_formatting.page_break_before if def_formatting.page_break_before is not None else False
    # Верт отступ перед абзацем
    space_before = formatting.space_before.pt if formatting.space_before is not None else \
        st_formatting.space_before.pt if st_formatting.space_before is not None else \
            def_formatting.space_before.pt if def_formatting.space_before is not None else 0.0
    # Верт отступ перед абзацем
    space_after = formatting.space_after.pt if formatting.space_after is not None else \
        st_formatting.space_after.pt if st_formatting.space_after is not None else \
            def_formatting.space_after.pt if def_formatting.space_after is not None else 0.0
    # Отступ слева
    left_indent = round(formatting.left_indent.cm, 2) if formatting.left_indent is not None else \
        round(st_formatting.left_indent.cm, 2) if st_formatting.left_indent is not None else \
            round(def_formatting.left_indent.cm, 2) if def_formatting.left_indent is not None else 0.0
    # Отступ справа
    right_indent = round(formatting.right_indent.cm, 2) if formatting.right_indent is not None else \
        round(st_formatting.right_indent.cm, 2) if st_formatting.right_indent is not None else \
            round(def_formatting.right_indent.cm, 2) if def_formatting.right_indent is not None else 0.0
    # Красная строка
    first_line_indent = round(formatting.first_line_indent.cm, 2) if formatting.first_line_indent is not None else \
        round(st_formatting.first_line_indent.cm, 2) if st_formatting.first_line_indent is not None else \
            round(def_formatting.first_line_indent.cm, 2) if def_formatting.first_line_indent is not None else 0.0
    # Межстрочный интервал
    line_spacing = formatting.line_spacing if formatting.line_spacing is not None else \
        st_formatting.line_spacing if st_formatting.line_spacing is not None else \
            def_formatting.line_spacing if def_formatting.line_spacing is not None else 1.0


    paragraph_stats = {
        "font_name": font_name,
        "font_size": font_size,
        "font_bald": font_bald,
        "font_italic": font_italic,
        "font_underline": font_underline,
        "font_color": font_color,
        "font_back_color": font_back_color,
        "alignment": alignment,
        "keep_with_next": keep_with_next,
        "page_break_before": page_break_before,
        "space_before": space_before,
        "space_after": space_after,
        "left_indent": left_indent,
        "right_indent": right_indent,
        "first_line_indent": first_line_indent,
        "line_spacing": line_spacing,
    }
    return paragraph_stats


# This function extracts the tables and paragraphs from the document object
def iter_block_items(parent):
    """
    Yield each paragraph and table child within *parent*, in document order.
    Each returned value is an instance of either Table or Paragraph. *parent*
    would most commonly be a reference to a main Document object, but
    also works for a _Cell object, which itself can contain paragraphs and tables.
    """
    if isinstance(parent, doctwo):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def parse_margins(document):
    sections = document.sections
    paragraph_stats = {}
    stats_to_compare = margins_check
    margins_comments = {}
    for section_cnt, section in enumerate(sections):
        paragraph_stats["top_margin"] = round(section.top_margin.cm, 2)
        paragraph_stats["bottom_margin"] = round(section.bottom_margin.cm, 2)
        paragraph_stats["left_margin"] = round(section.left_margin.cm, 2)
        paragraph_stats["right_margin"] = round(section.right_margin.cm, 2)
        paragraph_stats["orientation"] = section.orientation
        section_comments = get_error_comment(stats_to_compare, paragraph_stats)
        if len(section_comments) > 0:
            margins_comments[f"Раздел {section_cnt + 1}:"] = section_comments
    return margins_comments


def set_comment_name(name, prev_name, is_locked):
    if not is_locked:
        return name
    return prev_name


def parse_document(filename):
    document = Document(filename)
    comment_count = 0
    written_comments = []
    paragraph_to_comment = ""
    comment_to_send = set()
    comment_type = "System"
    image_name_check = 0
    first_paragraph_not_reached = True
    paragraphs = document.paragraphs
    for block in iter_block_items(document):
        next_paragraph_to_comment = ""
        next_comment_to_send = set()
        next_comment_type = "System"
        lock_comment_name = False
        stats_to_compare = {}
        run_comments = set()
        if 'text' in str(block):
            if first_paragraph_not_reached:
                first_paragraph_not_reached = False
                margin_comments = parse_margins(document)
                for section, comments in margin_comments.items():
                    block.add_comment(''.join(comments), author=section)
            for run in block.runs:
                xmlstr = str(run.element.xml)
                my_namespaces = dict([node for _, node in ET.iterparse(StringIO(xmlstr), events=['start-ns'])])
                root = ET.fromstring(xmlstr)
                if 'pic:pic' in xmlstr:  # Image run
                    for pic in root.findall('.//pic:pic', my_namespaces):
                        stats_to_compare = image_checklist
                        image_name_check = 2
                        paragraph_stats = get_run_properties(document, block, run)
                        run_comments = get_error_comment(stats_to_compare, paragraph_stats)
                        for comment in run_comments:
                            next_comment_to_send.add(comment)
                        if not lock_comment_name:
                            next_comment_type = set_comment_name("Рисунок", next_comment_type, lock_comment_name)
                        lock_comment_name = True
                        next_paragraph_to_comment = block
                        # print("Image:", block)
                else:  # Paragraph run
                    if run.text.strip() != "":

                        paragraph_stats = get_run_properties(document, block, run)
                        if block.style.name.startswith("Heading 1"):
                            paragraph_stats["is_list"] = is_list(block)
                            stats_to_compare = heading1_checklist
                            next_comment_type = set_comment_name("Заголовок 1", next_comment_type, lock_comment_name)
                        elif block.style.name.startswith("Heading 2"):
                            paragraph_stats["is_list"] = is_list(block)
                            stats_to_compare = heading2_checklist
                            next_comment_type = set_comment_name("Заголовок 2", next_comment_type, lock_comment_name)
                        elif block.style.name.startswith("Heading 3"):
                            paragraph_stats["is_list"] = is_list(block)
                            stats_to_compare = heading3_checklist
                            next_comment_type = set_comment_name("Заголовок 3", next_comment_type, lock_comment_name)
                        elif is_list(block):
                            stats_to_compare = list_checklist
                            next_comment_type = set_comment_name("Список", next_comment_type, lock_comment_name)
                        else:
                            stats_to_compare = text_checklist
                            next_comment_type = set_comment_name("Абзац", next_comment_type, lock_comment_name)
                        if enable_elem_check["enable_pic_title"] and image_name_check:
                            stats_to_compare = image_name_checklist
                            paragraph_stats["format_regex"] = block.text
                            next_comment_type = set_comment_name("Название рисунка", next_comment_type, lock_comment_name)
                            lock_comment_name = True
                        run_comments = get_error_comment(stats_to_compare, paragraph_stats)
                        for comment in run_comments:
                            next_comment_to_send.add(comment)
                        next_paragraph_to_comment = block
        elif 'table' in str(block):  # Table
            """Перепроверка названия таблица"""
            if enable_elem_check["paragraph_before_table"] and isinstance(paragraph_to_comment, Paragraph):
                comment_to_send.clear()
                stats_to_compare = table_name_checklist
                paragraph_stats["format_regex"] = paragraph_to_comment.text
                run_comments = get_error_comment(stats_to_compare, paragraph_stats)
                #print(run_comments)
                for comment in run_comments:
                    comment_to_send.add(comment)
                comment_type = "Название таблицы"
                #print('///', paragraph_to_comment.text, comment_to_send)
            #print("Table:", block)
            paragraph_chosen = False
            for row_count, row in enumerate(block.rows):
                col_count = 0
                for cell in row.cells:
                    if cell.text != "":
                        for par in cell.paragraphs:
                            if not paragraph_chosen:
                                next_paragraph_to_comment = par
                                paragraph_chosen = True
                            for run in par.runs:
                                paragraph_stats = get_run_properties(document, par, run)
                                vert_alignment = cell.vertical_alignment if cell.vertical_alignment else \
                                    WD_ALIGN_VERTICAL.TOP
                                paragraph_stats['vert_alignment'] = vert_alignment
                                if enable_elem_check["table_headings_top"] and row_count == 0:
                                    stats_to_compare = table_headings_checklist
                                elif enable_elem_check["table_headings_left"] and col_count == 0:
                                    stats_to_compare = table_headings_checklist
                                else:
                                    stats_to_compare = table_text_checklist

                                run_comments = get_error_comment(stats_to_compare, paragraph_stats)
                                # print(run_comments)
                                ''' print(par.text)
                                print(par.style.name)
                                print(paragraph_stats)
                                print(stats_to_compare)'''
                                # print(list(map(lambda k, v: v == stats_to_compare[k], paragraph_stats.keys(), paragraph_stats.values())))
                                for comment in run_comments:
                                    next_comment_to_send.add(comment)
                                next_comment_type = set_comment_name("Таблица", next_comment_type, lock_comment_name)
                    col_count += 1

        #print(paragraph_to_comment, comment_to_send)
        if isinstance(paragraph_to_comment, Paragraph) and len(comment_to_send) > 0:
            paragraph_to_comment.add_comment(''.join(comment_to_send), author=comment_type)
            comment_count += 1
            written_comments.append(comment_to_send)
        paragraph_to_comment = next_paragraph_to_comment
        comment_to_send = next_comment_to_send
        comment_type = next_comment_type
        image_name_check -= 1 if image_name_check > 0 else 0
    if isinstance(paragraph_to_comment, Paragraph) and len(comment_to_send) > 0:
        paragraph_to_comment.add_comment(''.join(comment_to_send), author=comment_type)
        comment_count += 1
        written_comments.append(comment_to_send)

    file, extension = os.path.splitext(filename)
    basename, extension = os.path.splitext(os.path.basename(filename))
    try:
        Path('./Results').mkdir(parents=True, exist_ok=False)
    except FileExistsError:
        pass
    count = 0
    while True:
        try:
            if count == 0:
                document.save(f"Results/{basename}_Проверенный{extension}")
            else:
                document.save(f"Results/{basename}_Проверенный_{count}{extension}")
            break
        except PermissionError:
            count += 1
    return written_comments


if __name__ == '__main__':

    # doc = Document("WordForTest.docx")

    '''for i in doc.paragraphs:
        print(is_list(i))'''
    # parse_document(doc)